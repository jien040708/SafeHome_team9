"""
SafeHomeCamera & CameraController Unit Test Document Generator with Test Function Names
두 클래스의 모든 메서드에 대한 단위 테스트 케이스 문서 생성 (테스트 함수 이름 포함)
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_heading('SafeHome Surveillance System Unit Test Cases', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============== SafeHomeCamera Class ==============
doc.add_heading('SafeHomeCamera Class', 1)

# SafeHomeCamera Header
header_table = doc.add_table(rows=1, cols=5)
header_table.style = 'Light Grid Accent 1'
header_cells = header_table.rows[0].cells
header_cells[0].text = 'Class'
header_cells[1].text = 'Method'
header_cells[2].text = 'Author'
header_cells[3].text = 'Date'
header_cells[4].text = 'Version'

info_row = header_table.add_row()
info_row.cells[0].text = 'SafeHomeCamera'
info_row.cells[1].text = 'All Methods'
info_row.cells[2].text = 'Jien Lee'
info_row.cells[3].text = datetime.now().strftime('%d.%m.%y')
info_row.cells[4].text = '1.0.1'

doc.add_paragraph()

# SafeHomeCamera Test Cases Table (7 columns)
safehome_table = doc.add_table(rows=1, cols=7)
safehome_table.style = 'Light Grid Accent 1'

# Header row
header_row = safehome_table.rows[0].cells
header_row[0].text = 'Test Case Description'
header_row[1].text = 'Method'
header_row[2].text = 'Input Specifications'
header_row[3].text = 'Expected Result'
header_row[4].text = 'Actual Result\n(Pass/Fail/Exception)'
header_row[5].text = 'Test Function Name'
header_row[6].text = 'Comment\n(including references)'

for cell in header_row:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

# SafeHomeCamera Test Cases with test function names
safehome_test_cases = [
    {'method': '__init__', 'desc': 'TC-001: Verifies that SafeHomeCamera initializes correctly with camera_id and optional location parameter.', 'input': 'Create SafeHomeCamera(camera_id=1, location=[100, 200]). Check attributes.', 'expected': 'Camera ID is 1, location is [100, 200], enabled is True, has_password is False.', 'actual': 'Pass: All attributes initialized correctly.', 'test_func': 'test_initialization_with_location', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:24-32'},
    {'method': '__init__', 'desc': 'TC-002: Verifies that SafeHomeCamera initializes with default location [0, 0] when location parameter is not provided.', 'input': 'Create SafeHomeCamera(camera_id=2) without location parameter. Call get_location().', 'expected': 'get_location() returns [0, 0].', 'actual': 'Pass: get_location() returns [0, 0].', 'test_func': 'test_initialization_without_location', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:34-38'},
    {'method': 'get_location', 'desc': 'TC-003: Verifies that get_location() returns the correct location coordinates and returns a copy.', 'input': 'Call get_location() on camera with location=[100, 200]. Modify returned list: location[0] = 999. Call get_location() again.', 'expected': 'First call returns [100, 200]. After modification, second call still returns [100, 200].', 'actual': 'Pass: Returns [100, 200] initially and after modification.', 'test_func': 'test_get_location', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:40-46'},
    {'method': 'set_location', 'desc': 'TC-004: Verifies that set_location() successfully updates the camera location when valid coordinates are provided.', 'input': 'Call set_location([300, 400]) on camera with initial location [100, 200]. Then call get_location().', 'expected': 'set_location() returns True, and get_location() returns [300, 400].', 'actual': 'Pass: set_location() returns True, and get_location() returns [300, 400].', 'test_func': 'test_set_location', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:48-52'},
    {'method': 'set_location', 'desc': 'TC-005: Verifies that set_location() returns False and preserves original location when invalid input is provided.', 'input': 'Call set_location([100]) (only x-coordinate) on camera with location [100, 200]. Call get_location().', 'expected': 'set_location() returns False, and get_location() returns [100, 200] (unchanged).', 'actual': 'Pass: set_location() returns False, and location remains [100, 200].', 'test_func': 'test_set_invalid_location', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:54-58'},
    {'method': 'get_id', 'desc': 'TC-006: Verifies that get_id() returns the correct camera ID.', 'input': 'Create camera with camera_id=1. Call get_id().', 'expected': 'get_id() returns 1.', 'actual': 'Pass: get_id() returns 1.', 'test_func': 'test_get_id', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:60-62'},
    {'method': 'set_id', 'desc': 'TC-007: Verifies that set_id() successfully updates the camera ID when a valid positive integer is provided.', 'input': 'Call set_id(5) on camera with ID 1. Call get_id().', 'expected': 'set_id() returns True, and get_id() returns 5.', 'actual': 'Pass: set_id() returns True, and get_id() returns 5.', 'test_func': 'test_set_id', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:64-68'},
    {'method': 'set_id', 'desc': 'TC-008: Verifies that set_id() returns False and preserves original ID when invalid ID (0 or negative) is provided.', 'input': 'Call set_id(0) and set_id(-1) on camera with ID 1. Call get_id() after each.', 'expected': 'Both set_id() calls return False, and get_id() returns 1 (unchanged).', 'actual': 'Pass: Both return False, ID remains 1.', 'test_func': 'test_set_invalid_id', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:70-77'},
    {'method': 'display_view', 'desc': 'TC-009: Verifies that display_view() returns a PIL Image when the camera is enabled.', 'input': 'Call display_view() on an enabled camera.', 'expected': 'display_view() returns a PIL Image.Image object (not None).', 'actual': 'Pass: Returns Image.Image object.', 'test_func': 'test_display_view', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:79-83'},
    {'method': 'display_view', 'desc': 'TC-010: Verifies that display_view() returns None when the camera is disabled.', 'input': 'Disable camera, then call display_view().', 'expected': 'display_view() returns None.', 'actual': 'Pass: Returns None when disabled.', 'test_func': 'test_display_view_disabled_camera', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:85-89'},
    {'method': 'zoom_in', 'desc': 'TC-011: Verifies that zoom_in() increases the zoom level by 1 when camera is enabled.', 'input': 'Get initial zoom setting, call zoom_in(), then get zoom setting again.', 'expected': 'zoom_in() returns True, and new zoom setting equals initial zoom + 1.', 'actual': 'Pass: Returns True, zoom increases by 1.', 'test_func': 'test_zoom_in', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:91-96'},
    {'method': 'zoom_in', 'desc': 'TC-012: Verifies that zoom_in() returns False when camera is disabled.', 'input': 'Disable camera, get initial zoom, call zoom_in(), check zoom again.', 'expected': 'zoom_in() returns False, and zoom setting remains unchanged.', 'actual': 'Pass: Returns False, zoom unchanged.', 'test_func': 'test_zoom_in_disabled_camera', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:109-115'},
    {'method': 'zoom_out', 'desc': 'TC-013: Verifies that zoom_out() decreases the zoom level by 1 when camera is enabled and zoom > 1.', 'input': 'Zoom in first, get zoom setting, call zoom_out(), check zoom again.', 'expected': 'zoom_out() returns True, and new zoom equals previous zoom - 1.', 'actual': 'Pass: Returns True, zoom decreases by 1.', 'test_func': 'test_zoom_out', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:117-123'},
    {'method': 'zoom_out', 'desc': 'TC-014: Verifies that zoom_out() returns False when zoom is already at minimum (1).', 'input': 'Call zoom_out() when zoom setting is 1.', 'expected': 'zoom_out() returns False, and zoom remains 1.', 'actual': 'Pass: Returns False, zoom stays at 1.', 'test_func': 'test_zoom_out_min', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:125-130'},
    {'method': 'pan_left', 'desc': 'TC-015: Verifies that pan_left() decreases pan angle by 5.0 degrees when camera is enabled.', 'input': 'Get initial pan angle, call pan_left(), check pan angle again.', 'expected': 'pan_left() returns True, and new pan angle equals initial - 5.0.', 'actual': 'Pass: Returns True, pan angle decreases by 5.0.', 'test_func': 'test_pan_left', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:141-146'},
    {'method': 'pan_right', 'desc': 'TC-016: Verifies that pan_right() increases pan angle by 5.0 degrees when camera is enabled.', 'input': 'Get initial pan angle, call pan_right(), check pan angle again.', 'expected': 'pan_right() returns True, and new pan angle equals initial + 5.0.', 'actual': 'Pass: Returns True, pan angle increases by 5.0.', 'test_func': 'test_pan_right', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:168-173'},
    {'method': 'get_password', 'desc': 'TC-017: Verifies that get_password() returns an empty string when no password is set.', 'input': 'Call get_password() on camera with no password set.', 'expected': 'get_password() returns empty string "".', 'actual': 'Pass: Returns empty string.', 'test_func': 'test_get_password_not_set', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:195-198'},
    {'method': 'set_password', 'desc': 'TC-018: Verifies that set_password() successfully sets a non-empty password.', 'input': 'Call set_password("mypassword123"), then call get_password() and has_password().', 'expected': 'set_password() returns True, get_password() returns "mypassword123", has_password() returns True.', 'actual': 'Pass: Password set successfully.', 'test_func': 'test_set_password', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:200-205'},
    {'method': 'set_password', 'desc': 'TC-019: Verifies that set_password() returns False when empty string is provided.', 'input': 'Call set_password(""), then check has_password().', 'expected': 'set_password() returns False, and has_password() returns False.', 'actual': 'Pass: Returns False, password not set.', 'test_func': 'test_set_empty_password', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:207-211'},
    {'method': 'is_enabled', 'desc': 'TC-020: Verifies that is_enabled() returns True for a newly created camera and False after disable().', 'input': 'Check is_enabled() initially, call disable(), check again, call enable(), check again.', 'expected': 'Initially True, after disable() returns False, after enable() returns True.', 'actual': 'Pass: Returns correct state at each step.', 'test_func': 'test_is_enabled', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:220-228'},
    {'method': 'enable', 'desc': 'TC-021: Verifies that enable() sets camera enabled state to True.', 'input': 'Disable camera, call enable(), then check is_enabled().', 'expected': 'is_enabled() returns True after enable() call.', 'actual': 'Pass: Returns True after enable().', 'test_func': 'test_enable', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:230-236'},
    {'method': 'disable', 'desc': 'TC-022: Verifies that disable() sets camera enabled state to False.', 'input': 'Call disable(), then check is_enabled().', 'expected': 'is_enabled() returns False after disable() call.', 'actual': 'Pass: Returns False after disable().', 'test_func': 'test_disable', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:238-243'},
    {'method': 'has_password', 'desc': 'TC-023: Verifies that has_password() returns False initially and True after setting a password.', 'input': 'Check has_password() initially, set password, check again.', 'expected': 'Initially False, after set_password() returns True.', 'actual': 'Pass: Returns correct state.', 'test_func': 'test_has_password', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:213-218'},
    {'method': 'save_info', 'desc': 'TC-024: Verifies that save_info() successfully saves camera information and returns True.', 'input': 'Call save_info() on a camera instance.', 'expected': 'save_info() returns True.', 'actual': 'Pass: Returns True.', 'test_func': 'test_save_info', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:245-248'},
    {'method': 'get_pan_angle', 'desc': 'TC-025: Verifies that get_pan_angle() returns the current pan angle from DeviceCamera.', 'input': 'Call get_pan_angle() on a camera, then pan_right(), call get_pan_angle() again.', 'expected': 'Initial angle is 0.0, after pan_right() angle is 5.0.', 'actual': 'Pass: Returns correct pan angles.', 'test_func': 'test_get_pan_angle', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:250-259'},
    {'method': 'get_zoom_setting', 'desc': 'TC-026: Verifies that get_zoom_setting() returns the current zoom level from DeviceCamera.', 'input': 'Call get_zoom_setting() initially, call zoom_in(), call get_zoom_setting() again.', 'expected': 'Initial zoom is 1, after zoom_in() zoom is 2.', 'actual': 'Pass: Returns correct zoom levels.', 'test_func': 'test_get_zoom_setting', 'comment': 'Reference: tests/surveillance_tests/test_safehome_camera.py:261-269'}
]

# Add SafeHomeCamera test cases
for tc in safehome_test_cases:
    row = safehome_table.add_row()
    row.cells[0].text = tc['desc']
    row.cells[1].text = tc['method']
    row.cells[2].text = tc['input']
    row.cells[3].text = tc['expected']
    row.cells[4].text = tc['actual']
    row.cells[5].text = tc['test_func']
    row.cells[6].text = tc['comment']

# Set column widths for SafeHomeCamera table
for col in safehome_table.columns:
    col.width = Inches(1.5)

# Set font size for SafeHomeCamera table
for row in safehome_table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)
            paragraph_format.space_after = Pt(6)
            for run in paragraph.runs:
                run.font.size = Pt(9)

# Page break
doc.add_page_break()

# ============== CameraController Class ==============
doc.add_heading('CameraController Class', 1)

# CameraController Header
header_table2 = doc.add_table(rows=1, cols=5)
header_table2.style = 'Light Grid Accent 1'
header_cells2 = header_table2.rows[0].cells
header_cells2[0].text = 'Class'
header_cells2[1].text = 'Method'
header_cells2[2].text = 'Author'
header_cells2[3].text = 'Date'
header_cells2[4].text = 'Version'

info_row2 = header_table2.add_row()
info_row2.cells[0].text = 'CameraController'
info_row2.cells[1].text = 'All Methods'
info_row2.cells[2].text = 'Jien Lee'
info_row2.cells[3].text = datetime.now().strftime('%d.%m.%y')
info_row2.cells[4].text = '1.0.1'

doc.add_paragraph()

# CameraController Test Cases Table (7 columns)
controller_table = doc.add_table(rows=1, cols=7)
controller_table.style = 'Light Grid Accent 1'

# Header row
header_row2 = controller_table.rows[0].cells
header_row2[0].text = 'Test Case Description'
header_row2[1].text = 'Method'
header_row2[2].text = 'Input Specifications'
header_row2[3].text = 'Expected Result'
header_row2[4].text = 'Actual Result\n(Pass/Fail/Exception)'
header_row2[5].text = 'Test Function Name'
header_row2[6].text = 'Comment\n(including references)'

for cell in header_row2:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

# CameraController Test Cases with test function names
controller_test_cases = [
    {'method': '__init__', 'desc': 'TC-001: Verifies that CameraController initializes with empty cameras dictionary, next_camera_id=1, and total_camera_number=0.', 'input': 'Create CameraController(), check _cameras, _next_camera_id, _total_camera_number.', 'expected': '_cameras is empty dict, _next_camera_id is 1, _total_camera_number is 0.', 'actual': 'Pass: All attributes initialized correctly.', 'test_func': 'test_initialization', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:25-31'},
    {'method': 'add_camera', 'desc': 'TC-002: Verifies that add_camera() successfully adds a new camera with specified coordinates and assigns sequential ID.', 'input': 'Call add_camera(100, 200), check _total_camera_number, _cameras, _next_camera_id.', 'expected': 'add_camera() returns True, _total_camera_number is 1, camera ID 1 exists in _cameras, _next_camera_id is 2.', 'actual': 'Pass: Camera added successfully with ID 1.', 'test_func': 'test_add_camera', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:33-43'},
    {'method': 'delete_camera', 'desc': 'TC-003: Verifies that delete_camera() successfully removes a camera and decreases total count.', 'input': 'Add two cameras, call delete_camera(1), check _total_camera_number and _cameras.', 'expected': 'delete_camera() returns True, _total_camera_number is 1, camera ID 1 not in _cameras, camera ID 2 still exists.', 'actual': 'Pass: Camera 1 deleted successfully.', 'test_func': 'test_delete_camera', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:58-68'},
    {'method': 'delete_camera', 'desc': 'TC-004: Verifies that delete_camera() returns False when attempting to delete non-existent camera.', 'input': 'Call delete_camera(999) on empty controller.', 'expected': 'delete_camera() returns False, _total_camera_number remains 0.', 'actual': 'Pass: Returns False for non-existent camera.', 'test_func': 'test_delete_nonexistent_camera', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:70-74'},
    {'method': 'enable_camera', 'desc': 'TC-005: Verifies that enable_camera() successfully enables a specific camera.', 'input': 'Add camera, disable it, call enable_camera(1), check camera.is_enabled().', 'expected': 'enable_camera() returns True, and camera.is_enabled() returns True.', 'actual': 'Pass: Camera enabled successfully.', 'test_func': 'test_enable_camera', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:76-83'},
    {'method': 'disable_camera', 'desc': 'TC-006: Verifies that disable_camera() successfully disables a specific camera.', 'input': 'Add camera, call disable_camera(1), check camera.is_enabled().', 'expected': 'disable_camera() returns True, and camera.is_enabled() returns False.', 'actual': 'Pass: Camera disabled successfully.', 'test_func': 'test_disable_camera', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:90-96'},
    {'method': 'enable_cameras', 'desc': 'TC-007: Verifies that enable_cameras() successfully enables multiple cameras from a list.', 'input': 'Add 3 cameras, disable all, call enable_cameras([1, 2, 3]), check each camera.is_enabled().', 'expected': 'enable_cameras() returns True, and all specified cameras are enabled.', 'actual': 'Pass: All cameras in list enabled successfully.', 'test_func': 'test_enable_cameras_list', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:103-115'},
    {'method': 'disable_cameras', 'desc': 'TC-008: Verifies that disable_cameras() successfully disables multiple cameras from a list.', 'input': 'Add 3 cameras, call disable_cameras([1, 2]), check each camera.is_enabled().', 'expected': 'disable_cameras() returns True, and specified cameras are disabled.', 'actual': 'Pass: Specified cameras disabled successfully.', 'test_func': 'test_disable_cameras_list', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:117-125'},
    {'method': 'enable_all_camera', 'desc': 'TC-009: Verifies that enable_all_camera() successfully enables all existing cameras.', 'input': 'Add 3 cameras, disable all, call enable_all_camera(), check all cameras.', 'expected': 'enable_all_camera() returns True, and all cameras are enabled.', 'actual': 'Pass: All cameras enabled successfully.', 'test_func': 'test_enable_all_cameras', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:127-137'},
    {'method': 'disable_all_camera', 'desc': 'TC-010: Verifies that disable_all_camera() successfully disables all existing cameras.', 'input': 'Add 3 cameras, call disable_all_camera(), check all cameras.', 'expected': 'disable_all_camera() returns True, and all cameras are disabled.', 'actual': 'Pass: All cameras disabled successfully.', 'test_func': 'test_disable_all_cameras', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:139-147'},
    {'method': 'control_single_camera', 'desc': 'TC-011: Verifies that control_single_camera() executes pan_left when control_id=0 on enabled camera.', 'input': 'Add camera, call control_single_camera(1, 0).', 'expected': 'control_single_camera() returns True.', 'actual': 'Pass: Pan left executed successfully.', 'test_func': 'test_control_single_camera_pan_left', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller_coverage.py:85-90'},
    {'method': 'control_single_camera', 'desc': 'TC-012: Verifies that control_single_camera() executes pan_right when control_id=1 on enabled camera.', 'input': 'Add camera, call control_single_camera(1, 1).', 'expected': 'control_single_camera() returns True.', 'actual': 'Pass: Pan right executed successfully.', 'test_func': 'test_control_single_camera_pan_right', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller_coverage.py:93-98'},
    {'method': 'control_single_camera', 'desc': 'TC-013: Verifies that control_single_camera() executes zoom_in when control_id=2 on enabled camera.', 'input': 'Add camera, call control_single_camera(1, 2).', 'expected': 'control_single_camera() returns True.', 'actual': 'Pass: Zoom in executed successfully.', 'test_func': 'test_control_single_camera_zoom_in', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller_coverage.py:101-106'},
    {'method': 'control_single_camera', 'desc': 'TC-014: Verifies that control_single_camera() executes zoom_out when control_id=3 on enabled camera.', 'input': 'Add camera, zoom in first, call control_single_camera(1, 3).', 'expected': 'control_single_camera() returns True.', 'actual': 'Pass: Zoom out executed successfully.', 'test_func': 'test_control_single_camera_zoom_out', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller_coverage.py:109-115'},
    {'method': 'control_single_camera', 'desc': 'TC-015: Verifies that control_single_camera() returns False for non-existent camera.', 'input': 'Call control_single_camera(999, 0).', 'expected': 'control_single_camera() returns False.', 'actual': 'Pass: Returns False for non-existent camera.', 'test_func': 'test_control_nonexistent_camera', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:180-183'},
    {'method': 'control_single_camera', 'desc': 'TC-016: Verifies that control_single_camera() returns False when camera is disabled.', 'input': 'Add camera, disable it, call control_single_camera(1, 0).', 'expected': 'control_single_camera() returns False.', 'actual': 'Pass: Returns False for disabled camera.', 'test_func': 'test_control_disabled_camera', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:172-178'},
    {'method': 'display_thumbnail_view', 'desc': 'TC-017: Verifies that display_thumbnail_view() returns a composite thumbnail image when cameras exist.', 'input': 'Add 3 cameras, call display_thumbnail_view().', 'expected': 'display_thumbnail_view() returns a PIL Image object with thumbnails of all cameras.', 'actual': 'Pass: Returns composite thumbnail image.', 'test_func': 'test_display_thumbnail_view_with_multiple_cameras', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller_coverage.py:118-126'},
    {'method': 'display_thumbnail_view', 'desc': 'TC-018: Verifies that display_thumbnail_view() returns None when no cameras exist.', 'input': 'Call display_thumbnail_view() on empty controller.', 'expected': 'display_thumbnail_view() returns None.', 'actual': 'Pass: Returns None when no cameras.', 'test_func': 'test_display_thumbnail_view_no_cameras', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:192-195'},
    {'method': 'display_single_view', 'desc': 'TC-019: Verifies that display_single_view() returns camera image for enabled camera.', 'input': 'Add camera, call display_single_view(1).', 'expected': 'display_single_view() returns PIL Image object.', 'actual': 'Pass: Returns camera image.', 'test_func': 'test_display_single_view', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:197-203'},
    {'method': 'display_single_view', 'desc': 'TC-020: Verifies that display_single_view() returns None for disabled camera.', 'input': 'Add camera, disable it, call display_single_view(1).', 'expected': 'display_single_view() returns None.', 'actual': 'Pass: Returns None for disabled camera.', 'test_func': 'test_display_single_view_disabled_camera', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:205-211'},
    {'method': 'get_all_camera_info', 'desc': 'TC-021: Verifies that get_all_camera_info() returns list of camera information arrays [id, x, y, enabled, total_count].', 'input': 'Add 2 cameras at different locations, call get_all_camera_info().', 'expected': 'Returns list with 2 arrays, each containing [camera_id, x, y, enabled_flag, total_count].', 'actual': 'Pass: Returns correct camera information list.', 'test_func': 'test_get_all_camera_info', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:218-230'},
    {'method': 'set_camera_password', 'desc': 'TC-022: Verifies that set_camera_password() successfully sets password for specified camera.', 'input': 'Add camera, call set_camera_password(1, "test123"), then validate_camera_password(1, "test123").', 'expected': 'validate_camera_password() returns 0 (success).', 'actual': 'Pass: Password set and validated successfully.', 'test_func': 'test_set_camera_password', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:237-243'},
    {'method': 'validate_camera_password', 'desc': 'TC-023: Verifies that validate_camera_password() returns 0 when correct password is provided.', 'input': 'Add camera, set password "test123", call validate_camera_password(1, "test123").', 'expected': 'validate_camera_password() returns 0 (success).', 'actual': 'Pass: Returns 0 for correct password.', 'test_func': 'test_validate_camera_password_success', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:251-257'},
    {'method': 'validate_camera_password', 'desc': 'TC-024: Verifies that validate_camera_password() returns 1 when incorrect password is provided.', 'input': 'Add camera, set password "test123", call validate_camera_password(1, "wrong").', 'expected': 'validate_camera_password() returns 1 (incorrect password).', 'actual': 'Pass: Returns 1 for incorrect password.', 'test_func': 'test_validate_camera_password_failure', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:259-265'},
    {'method': 'validate_camera_password', 'desc': 'TC-025: Verifies that validate_camera_password() returns -1 when camera does not exist.', 'input': 'Call validate_camera_password(999, "any").', 'expected': 'validate_camera_password() returns -1 (camera not found).', 'actual': 'Pass: Returns -1 for non-existent camera.', 'test_func': 'test_validate_camera_password_nonexistent_camera', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:267-270'},
    {'method': 'validate_camera_password', 'desc': 'TC-026: Verifies that validate_camera_password() returns -2 when camera has no password set.', 'input': 'Add camera without setting password, call validate_camera_password(1, "any").', 'expected': 'validate_camera_password() returns -2 (no password set).', 'actual': 'Pass: Returns -2 when no password set.', 'test_func': 'test_validate_camera_password_not_set', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:272-277'},
    {'method': 'delete_camera_password', 'desc': 'TC-027: Verifies that delete_camera_password() successfully removes password for specified camera.', 'input': 'Add camera, set password, call delete_camera_password(1), then validate_camera_password(1, "old").', 'expected': 'delete_camera_password() returns 0, and validate returns -2 (no password).', 'actual': 'Pass: Password deleted successfully.', 'test_func': 'N/A (private method)', 'comment': 'This is a private method, tested indirectly through camera deletion.'},
    {'method': 'trigger_security_event', 'desc': 'TC-028: Verifies that trigger_security_event() captures images from all enabled cameras.', 'input': 'Add 3 cameras, enable all, call trigger_security_event("INTRUSION_ALARM").', 'expected': 'Images are captured from all 3 enabled cameras.', 'actual': 'Pass: All enabled cameras triggered.', 'test_func': 'test_trigger_security_event_with_enabled_cameras', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller_coverage.py:27-37'},
    {'method': 'trigger_security_event', 'desc': 'TC-029: Verifies that trigger_security_event() skips disabled cameras.', 'input': 'Add 2 cameras, disable one, call trigger_security_event("PANIC").', 'expected': 'Only enabled camera captures image.', 'actual': 'Pass: Disabled camera skipped.', 'test_func': 'test_trigger_security_event_with_disabled_camera', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller_coverage.py:40-52'},
    {'method': 'trigger_security_event', 'desc': 'TC-030: Verifies that trigger_security_event() handles exception from camera gracefully.', 'input': 'Add camera, simulate exception in camera.display_view(), call trigger_security_event("TEST").', 'expected': 'Process continues without crashing, exception is logged.', 'actual': 'Pass: Exception handled gracefully.', 'test_func': 'test_trigger_security_event_camera_exception', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller_coverage.py:63-71'},
    {'method': 'get_camera_count', 'desc': 'TC-031: Verifies that get_camera_count() returns the total number of cameras.', 'input': 'Add 3 cameras, call get_camera_count().', 'expected': 'get_camera_count() returns 3.', 'actual': 'Pass: Returns correct count.', 'test_func': 'test_get_camera_count', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:288-299'},
    {'method': 'get_camera', 'desc': 'TC-032: Verifies that get_camera() returns SafeHomeCamera object for existing camera.', 'input': 'Add camera, call get_camera(1).', 'expected': 'get_camera() returns SafeHomeCamera instance with ID 1.', 'actual': 'Pass: Returns camera object.', 'test_func': 'test_get_camera', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:301-308'},
    {'method': 'get_camera', 'desc': 'TC-033: Verifies that get_camera() returns None for non-existent camera.', 'input': 'Call get_camera(999) on empty controller.', 'expected': 'get_camera() returns None.', 'actual': 'Pass: Returns None for non-existent camera.', 'test_func': 'test_get_nonexistent_camera', 'comment': 'Reference: tests/surveillance_tests/test_camera_controller.py:310-313'}
]

# Add CameraController test cases
for tc in controller_test_cases:
    row = controller_table.add_row()
    row.cells[0].text = tc['desc']
    row.cells[1].text = tc['method']
    row.cells[2].text = tc['input']
    row.cells[3].text = tc['expected']
    row.cells[4].text = tc['actual']
    row.cells[5].text = tc['test_func']
    row.cells[6].text = tc['comment']

# Set column widths for CameraController table
for col in controller_table.columns:
    col.width = Inches(1.5)

# Set font size for CameraController table
for row in controller_table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)
            paragraph_format.space_after = Pt(6)
            for run in paragraph.runs:
                run.font.size = Pt(9)

# Save document
doc.save('Surveillance_System_Unit_Tests_With_Func_Names.docx')
print("Document created: Surveillance_System_Unit_Tests_With_Func_Names.docx")
print(f"Total test cases: SafeHomeCamera ({len(safehome_test_cases)}), CameraController ({len(controller_test_cases)})")

